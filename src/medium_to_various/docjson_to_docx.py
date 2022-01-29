from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches, Pt, RGBColor
from utils import jsonx

from medium_to_various.remote_file_utils import get_local_file

DEFAULT_IMAGE_WIDTH = 3
DEFAULT_FONT_NAME = 'Georgia'


def _build_styles(document):
    style = document.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = DEFAULT_FONT_NAME

    style = document.styles['Quote']
    style.font.color.rgb = RGBColor(128, 0, 0)
    style.font.name = DEFAULT_FONT_NAME
    style.paragraph_format.left_indent = Inches(0.25)
    style.paragraph_format.right_indent = Inches(0.25)

    style = document.styles['Caption']
    style.font.color.rgb = RGBColor(192, 0, 0)
    style.font.name = DEFAULT_FONT_NAME
    style.paragraph_format.left_indent = Inches(0.25)
    style.paragraph_format.right_indent = Inches(0.25)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.font.size = Pt(10)
    style.font.bold = False
    style.font.italic = True

    style = document.styles.add_style('New Heading 0', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(20)
    style.font.name = DEFAULT_FONT_NAME
    style.font.color.rgb = RGBColor(128, 128, 128)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = document.styles.add_style('New Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(15)
    style.font.name = DEFAULT_FONT_NAME
    style.font.color.rgb = RGBColor(96, 96, 96)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = document.styles.add_style('New Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(12)
    style.font.name = DEFAULT_FONT_NAME
    style.font.color.rgb = RGBColor(64, 64, 64)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = document.styles.add_style('New Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(10)
    style.font.name = DEFAULT_FONT_NAME
    style.font.color.rgb = RGBColor(32, 32, 32)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def docjson_to_docx(docjson_file, docx_file):
    docjson = jsonx.read(docjson_file)
    document = Document()
    _build_styles(document)
    for d in docjson:
        tag, text = d['tag'], d.get('text')
        if tag == 'title':
            document.add_paragraph(text, style='New Heading 0')
        elif tag == 'h1':
            document.add_paragraph(text, style='New Heading 1')
        elif tag == 'h2':
            document.add_paragraph(text, style='New Heading 2')
        elif tag == 'h3':
            document.add_paragraph(text, style='New Heading 3')

        elif tag == 'p':
            document.add_paragraph(text)
        elif tag == 'em':
            p = document.add_paragraph()
            p.add_run(text).italic = True
        elif tag == 'figcaption':
            document.add_paragraph(text, style='Caption')
        elif tag == 'blockquote':
            document.add_paragraph(text, style='Quote')

        elif tag == 'li':
            document.add_paragraph(text, style='List Bullet')
        elif tag == 'pre':
            document.add_paragraph(text)
        elif tag == 'img':
            url = d['src']
            local_file = get_local_file(url)
            try:
                document.add_picture(
                    local_file, width=Inches(DEFAULT_IMAGE_WIDTH)
                )
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except UnrecognizedImageError:
                print(f'Could not add image: "{url}"')

        elif tag == 'ellipsis':
            last_paragraph = document.add_paragraph('...')
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif tag == 'time':
            p = document.add_paragraph()
            text_expanded = f'''
Colombo, Sri Lanka,
{text}
            '''
            p.add_run(text_expanded).italic = True
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            document.add_page_break()

    document.save(docx_file)
    print(f'Wrote {docx_file}')
